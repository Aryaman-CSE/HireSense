from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
import pdfplumber
import docx
import pytesseract
from PIL import Image
import io
import os
from resume_parser import parse_resume

pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

app = FastAPI(title="AI Resume Parser")

SUPPORTED_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/msword": "doc",
    "image/png": "image",
    "image/jpeg": "image",
    "image/jpg": "image",
    "text/plain": "txt",
}

def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted + "\n"
    return text.strip()

def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = docx.Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)

def extract_text_from_image(file_bytes: bytes) -> str:
    image = Image.open(io.BytesIO(file_bytes))
    text = pytesseract.image_to_string(image)
    return text.strip()

def extract_text_from_txt(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore").strip()

def get_file_type(filename: str, content_type: str) -> str:
    ext = os.path.splitext(filename)[-1].lower()
    ext_map = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "doc",
        ".png": "image",
        ".jpg": "image",
        ".jpeg": "image",
        ".txt": "txt",
    }
    return ext_map.get(ext) or SUPPORTED_TYPES.get(content_type)

@app.post("/parse-resume")
async def parse_resume_api(file: UploadFile = File(...)):
    file_bytes = await file.read()
    file_type = get_file_type(file.filename or "", file.content_type or "")

    if file_type is None:
        raise HTTPException(
            status_code=415,
            detail="Unsupported file type. Accepted: PDF, DOCX, DOC, PNG, JPG, TXT"
        )

    if file_type == "pdf":
        text = extract_text_from_pdf(file_bytes)
    elif file_type in ("docx", "doc"):
        text = extract_text_from_docx(file_bytes)
    elif file_type == "image":
        text = extract_text_from_image(file_bytes)
    elif file_type == "txt":
        text = extract_text_from_txt(file_bytes)
    else:
        raise HTTPException(status_code=415, detail="Unsupported file format.")

    if not text:
        raise HTTPException(
            status_code=422,
            detail="Could not extract any text from the uploaded file."
        )

    result = parse_resume(text)
    result["source_format"] = file_type
    result["char_count"] = len(text)

    return JSONResponse(content=result)

@app.get("/health")
async def health():
    return {"status": "ok", "supported_formats": list(SUPPORTED_TYPES.values())}